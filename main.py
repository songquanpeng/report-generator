import datetime
import json
import os
import random
from docx import Document
import locale
import pywintypes, win32file, win32con

locale.setlocale(locale.LC_CTYPE, 'chinese')


def load_config(filename="./config.json"):
    with open(filename, encoding="utf-8") as f:
        config = json.load(f)
    if config["student_name"] == "张三":
        print("警告：似乎你忘记修改了 config.json 文件。")
    return config


def date_generator(week_num=8, working_day_num=5):
    date = datetime.datetime.strptime("08-31", "%m-%d")
    for week in range(week_num):
        for working_day in range(working_day_num):
            yield date.strftime("%m {m} %d {d}").format(m='月', d='日')
            if working_day == working_day_num - 1:
                delta = datetime.timedelta(days=3)
            else:
                delta = datetime.timedelta(days=1)
            date += delta


def generate_content_list(config, num):
    content = {"student_name": config["student_name"],
               "student_id": config["student_id"],
               "lab_name": config["lab_name"],
               "major": config["major"],
               "tutor_name": config["tutor_name"],
               "project_name": config["project_name"],
               "todo_task_1": "",
               "todo_task_2": "",
               "progress_1": "",
               "progress_2": "",
               "tomorrow_plan_1": "",
               "tomorrow_plan_2": "",
               "others_1": "",
               "others_2": ""
               }
    content_list = [content.copy() for _ in range(num)]
    paper_list = config["paper_list"]
    paper_list = ["阅读论文 " + paper for paper in paper_list]
    experiment_list = config["experiment_list"]
    experiment_list = ["进行实验" + experiment for experiment in experiment_list]
    task_list = config["task_list"]
    task_list.extend(paper_list)
    task_list.extend(experiment_list)
    others_list = config["others_list"]
    if len(task_list) < num:
        repeat_num = num - len(task_list)
        for _ in range(repeat_num):
            random_index = random.randint(0, len(task_list) - 1)
            task_list.insert(random_index, task_list[random_index])
    assert len(task_list) >= num
    course_list = []
    if config["machine_learning_course"]:
        course_list.extend(config["ml_course_list"])
    if config["deep_learning_course"]:
        course_list.extend(config["dl_course_list"])
    course_list = ["学习课程 " + course for course in course_list]
    for i in range(len(content_list)):
        if i < len(course_list):
            content_list[i]["todo_task_1"] = "完成相关课程的学习"
            content_list[i]["todo_task_2"] = "读相关的论文"
            content_list[i]["progress_2"] = course_list[i]
            if i + 1 != len(course_list):
                content_list[i]["tomorrow_plan_2"] = course_list[i + 1]
        else:
            content_list[i]["todo_task_1"] = "完成所要求的实验"
            content_list[i]["todo_task_2"] = "读相关的论文"
        content_list[i]["progress_1"] = task_list[i]
        if i + 1 != len(task_list):
            content_list[i]["tomorrow_plan_1"] = task_list[i + 1]
        if random.random() > 0.5:
            content_list[i]["others_1"] = random.choice(others_list)

    return content_list


def update_file_time(filename, date):
    new_datetime = datetime.datetime.now()
    new_datetime = new_datetime.replace(month=int(date[:2]), day=int(date[5:7]), hour=21, minute=random.randint(1, 59))
    timestamp = datetime.datetime.timestamp(new_datetime)
    os.utime(filename, (timestamp, timestamp))
    if os.name == "nt":
        win_time = pywintypes.Time(new_datetime)
        win_file = win32file.CreateFile(
            filename, win32con.GENERIC_WRITE,
            win32con.FILE_SHARE_READ | win32con.FILE_SHARE_WRITE | win32con.FILE_SHARE_DELETE,
            None, win32con.OPEN_EXISTING,
            win32con.FILE_ATTRIBUTE_NORMAL, None)

        win32file.SetFileTime(win_file, win_time, None, None)
        win_file.close()


def generate_log(content):
    date = content["date"]
    print(f"Generating log for day: {date}...", end=" ")
    source_filename = f"./template.docx"
    target_filename = f"./generated/{date}.docx"
    doc = Document(source_filename)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:  # Keep the style
                        for key in content.keys():
                            if key in run.text:
                                run.text = run.text.replace(key, content[key])
                                break
    doc.save(target_filename)
    update_file_time(target_filename, date)
    print("Done.")


def main():
    if not os.path.exists("generated"):
        os.makedirs("generated")
    cfg = load_config()
    week_num = 8
    working_day_num = 5
    date = date_generator(week_num, working_day_num)
    content_list = generate_content_list(cfg, week_num * working_day_num)
    for content in content_list:
        current_date = next(date)
        content["date"] = current_date
        generate_log(content)


if __name__ == '__main__':
    main()
